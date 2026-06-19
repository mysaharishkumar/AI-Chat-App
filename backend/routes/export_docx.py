from fastapi import APIRouter
from fastapi.responses import FileResponse

from docx import Document

from database.messages import messages

router = APIRouter(
    prefix="/export-docx",
    tags=["Export DOCX"]
)

@router.get("/{user_id}/{thread_id}")
async def export_docx(
    user_id: str,
    thread_id: str
):

    filename = (
        f"chat_{thread_id}.docx"
    )

    doc = Document()

    doc.add_heading(
        "AI Chat Export",
        level=1
    )

    async for msg in messages.find(
        {
            "user_id": user_id,
            "thread_id": thread_id
        }
    ):

        doc.add_paragraph(
            f"You: {msg['user_message']}"
        )

        doc.add_paragraph(
            f"AI: {msg['ai_response']}"
        )

        doc.add_paragraph(
            "----------------"
        )

    doc.save(
        filename
    )

    return FileResponse(
        filename,
        filename=filename
    )